from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

doc = Document()

def center_text(text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

def bold_paragraph(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(ns.qn('w:val'), 'single')
        border.set(ns.qn('w:sz'), '4')
        border.set(ns.qn('w:space'), '0')
        border.set(ns.qn('w:color'), '000000')
        borders.append(border)
    tcPr.append(borders)

center_text("FORM ‘A’", 12, True) #Header
center_text("MEDIATION APPLICATION FORM", 14, True)
center_text("[REFER RULE 3(1)]", 10)
center_text("Mumbai District Legal Services Authority", 11)
center_text("City Civil Court, Mumbai", 11)

doc.add_paragraph("\n")

table = doc.add_table(rows=1, cols=3)#table
table.style = "Table Grid"

hdr_cells = table.rows[0].cells
hdr_cells[0].text = ""
hdr_cells[1].text = "DETAILS OF PARTIES:"
hdr_cells[2].text = ""

for cell in hdr_cells:
    set_cell_border(cell)
    for p in cell.paragraphs:
        p.runs[0].bold = True

def add_row(col1, col2, col3=""):
    row = table.add_row().cells
    row[0].text = col1
    row[1].text = col2
    row[2].text = col3
    for cell in row:
        set_cell_border(cell)

add_row("1", "Name of Applicant", "{{client_name}}")
add_row("", "Address and contact details of Applicant", "")
add_row("1", "REGISTERED ADDRESS:\n{{branch_address}}\n\nCORRESPONDENCE BRANCH ADDRESS:\n{{branch_address}}", "")
add_row("", "Telephone No.", "{{mobile}}")
add_row("", "Mobile No.", "")
add_row("", "Email ID", "info@kslegal.co.in")

add_row("2", "Name, Address and Contact details of Opposite Party", "")
add_row("", "Name", "{{customer_name}}")
add_row("", "REGISTERED ADDRESS:\n{% if address1 and address1 != '' %}{{address1}}{% else %}____________{% endif %}", "")
add_row("", "CORRESPONDENCE ADDRESS:\n{% if address1 and address1 != '' %}{{address1}}{% else %}____________{% endif %}", "")
add_row("", "Telephone No.", "")
add_row("", "Mobile No.", "")
add_row("", "Email ID", "")


row = table.add_row().cells#final section
row[0].merge(row[1]).merge(row[2])
row[0].text = "DETAILS OF DISPUTE:"
for cell in row:
    set_cell_border(cell)
    cell.paragraphs[0].runs[0].bold = True

row = table.add_row().cells
row[0].merge(row[1]).merge(row[2])
row[0].text = "THE COMM. COURTS (PRE-INSTITUTION ……… SETTLEMENT) RULES, 2018"
for cell in row:
    set_cell_border(cell)
    cell.paragraphs[0].runs[0].bold = True

row = table.add_row().cells
row[0].merge(row[1]).merge(row[2])
row[0].text = "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):"
for cell in row:
    set_cell_border(cell)

doc.save("Mediation_Application_Form.docx")

